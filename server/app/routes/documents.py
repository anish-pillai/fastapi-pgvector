from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from uuid import UUID
import datetime
import hashlib
import os
from openai import OpenAI

from app.core.database import get_db
from app.core.config import settings
from app.models.models import Document
from app.schemas import DocumentResponse

router = APIRouter()


@router.post("/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    user_id: UUID = Form(...),
    db: Session = Depends(get_db)
):
    # Read file content
    content = await file.read()
    text = None
    # Detect and extract text based on file type
    if file.content_type == "text/plain" or file.filename.lower().endswith(".txt"):
        try:
            text = content.decode("utf-8")
        except Exception:
            raise HTTPException(status_code=400, detail="Could not decode text file as UTF-8.")
    elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file.filename.lower().endswith(".docx"):
        try:
            from docx import Document as DocxDocument
            from io import BytesIO
            docx = DocxDocument(BytesIO(content))
            text = "\n".join([p.text for p in docx.paragraphs])
        except Exception:
            raise HTTPException(status_code=400, detail="Could not extract text from Word document.")
    elif file.content_type == "application/pdf" or file.filename.lower().endswith(".pdf"):
        try:
            from PyPDF2 import PdfReader
            from io import BytesIO
            reader = PdfReader(BytesIO(content))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            raise HTTPException(status_code=400, detail="Could not extract text from PDF document.")
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {file.content_type}")

    if not text or not text.strip():
        raise HTTPException(status_code=400, detail="No extractable text found in file.")

    # Generate embeddings using OpenAI
    try:
        client = OpenAI(api_key=settings.OPENAI_API_KEY)
        embedding_response = client.embeddings.create(
            input=text,
            model="text-embedding-ada-002"
        )
        embedding = embedding_response.data[0].embedding
        n_tokens = embedding_response.usage.total_tokens if hasattr(embedding_response, 'usage') else len(text.split())
        cost = n_tokens / 1000 * 0.0001
    except Exception as e:
        print("Embedding error:", e)
        embedding = [0.0] * 1536
        cost = 0.0

    # Enrich metadata
    current_time = datetime.datetime.now(datetime.timezone.utc)
    metadata = {
        "filename": file.filename,
        "content_type": file.content_type,
        "size_bytes": len(content),
        "word_count": len(text.split()),
        "line_count": text.count("\n") + 1,
        "sha256": hashlib.sha256(content).hexdigest(),
        "upload_time": current_time.isoformat(),
        "upload_time_epoch": int(current_time.timestamp()),
        "embedding_model": "text-embedding-ada-002",
        "embedding_cost_usd": cost,
        "embedding_dim": len(embedding),
        "embedding_tokens": n_tokens if 'n_tokens' in locals() else len(text.split()),
    }

    db_doc = Document(
        filename=file.filename,
        doc_metadata=metadata,
        content=text,
        embedding=embedding,
        user_id=user_id
    )
    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)
    return db_doc


@router.get("/{doc_id}", response_model=DocumentResponse)
def get_document(doc_id: UUID, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc